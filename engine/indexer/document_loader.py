from typing import List, Dict, Optional
import os
from pathlib import Path
from pypdf import PdfReader
import docx
import pandas as pd
import markdown
from bs4 import BeautifulSoup
import tiktoken
from ..web.apiconfig import config

class DocumentLoader:
    def __init__(self, base_dir: str = "/Users/bojieli/pyproject/llm-search/knowledge_base/docs",
                 max_tokens_per_chunk: int = 1000):
        self.base_dir = Path(base_dir)
        self.max_tokens_per_chunk = max_tokens_per_chunk
        
        # 初始化 tokenizer
        try:
            self.tokenizer = tiktoken.encoding_for_model(config.api.embedding["model"])
        except Exception as e:
            print(f"初始化 tokenizer 失败: {e}")
            self.tokenizer = tiktoken.get_encoding("cl100k_base")

    def count_tokens(self, text: str) -> int:
        """计算文本的 token 数量"""
        return len(self.tokenizer.encode(text))

    def split_text(self, text: str, metadata: Dict) -> List[Dict]:
        """将文本分割成适当大小的块"""
        chunks = []
        current_chunk = ""
        current_tokens = 0
        
        # 按句子分割文本
        sentences = text.replace('\n', ' ').split('。')
        
        for sentence in sentences:
            sentence = sentence.strip() + '。'
            sentence_tokens = self.count_tokens(sentence)
            
            if current_tokens + sentence_tokens > self.max_tokens_per_chunk:
                if current_chunk:
                    chunks.append({
                        'content': current_chunk,
                        'token_count': current_tokens,
                        **metadata
                    })
                current_chunk = sentence
                current_tokens = sentence_tokens
            else:
                current_chunk += sentence
                current_tokens += sentence_tokens
        
        if current_chunk:
            chunks.append({
                'content': current_chunk,
                'token_count': current_tokens,
                **metadata
            })
        
        return chunks

    def _load_pdf(self, file_path: Path) -> List[Dict]:
        """加载 PDF 文件"""
        documents = []
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    metadata = {
                        'source': str(file_path),
                        'page': i + 1
                    }
                    documents.extend(self.split_text(text, metadata))
        return documents

    def _load_docx(self, file_path: Path) -> List[Dict]:
        """加载 Word 文档"""
        doc = docx.Document(file_path)
        documents = []
        current_text = ""
        current_para = 1
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                current_text += paragraph.text + "\n"
                if self.count_tokens(current_text) >= self.max_tokens_per_chunk:
                    metadata = {
                        'source': str(file_path),
                        'paragraph_range': f"{current_para}-{i+1}"
                    }
                    documents.extend(self.split_text(current_text, metadata))
                    current_text = ""
                    current_para = i + 2
        
        if current_text:
            metadata = {
                'source': str(file_path),
                'paragraph_range': f"{current_para}-{len(doc.paragraphs)}"
            }
            documents.extend(self.split_text(current_text, metadata))
        
        return documents

    def _load_excel(self, file_path: Path) -> List[Dict]:
        """加载 Excel 文件"""
        documents = []
        df = pd.read_excel(file_path)
        for sheet_name, sheet_df in df.items():
            content = sheet_df.to_string()
            if content.strip():
                metadata = {
                    'source': str(file_path),
                    'sheet': sheet_name
                }
                documents.extend(self.split_text(content, metadata))
        return documents

    def _load_markdown(self, file_path: Path) -> List[Dict]:
        """加载 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_text = file.read()
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            metadata = {
                'source': str(file_path)
            }
            return self.split_text(text, metadata)